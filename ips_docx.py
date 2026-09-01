"""
ips_docx.py

Parses an IPS-parameters Word document into the same dict shape used
everywhere else in this project (see sample_data/sample_ips.json).

This only understands one convention, matching IPS_Template.docx:
  - "Label: value" paragraphs for the scalar fields (client name, risk
    tolerance, time horizon, liquidity reserve, single-position limit)
  - a "Constraints" heading followed by a bullet (or "- "/"* ") list
  - a "Target Asset Allocation Bands" heading followed by a 3-column
    table: Asset Class | Min % | Max %

It is deliberately a template-following parser, not a general-purpose
document reader -- there's no AI step involved, so it can only find
what it's told to look for.
"""

from __future__ import annotations

from typing import Any

from docx import Document

_LABELS = {
    "client name": "client_name",
    "risk tolerance": "risk_tolerance",
    "time horizon (years)": "time_horizon_years",
    "time horizon": "time_horizon_years",
    "liquidity reserve requirement (%)": "liquidity_reserve_pct",
    "liquidity reserve requirement": "liquidity_reserve_pct",
    "liquidity reserve (%)": "liquidity_reserve_pct",
    "single-position limit (%)": "single_position_limit_pct",
    "single position limit (%)": "single_position_limit_pct",
    "single-position limit": "single_position_limit_pct",
    "single position limit": "single_position_limit_pct",
}

_NUMERIC_FIELDS = {"time_horizon_years", "liquidity_reserve_pct", "single_position_limit_pct"}


class IpsDocxParseError(ValueError):
    pass


def _clean_number(raw: str) -> float:
    return float(raw.replace("%", "").replace(",", "").strip())


def parse_ips_docx(file) -> dict[str, Any]:
    doc = Document(file)

    ips: dict[str, Any] = {
        "client_name": "",
        "risk_tolerance": "Moderate",
        "time_horizon_years": 10,
        "liquidity_reserve_pct": 5.0,
        "single_position_limit_pct": 10.0,
        "constraints": [],
        "allocation_targets": {},
    }

    in_constraints = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        low = text.lower()

        if low.startswith("constraints"):
            in_constraints = True
            continue
        if low.startswith("target asset allocation") or low.startswith("allocation bands"):
            in_constraints = False
            continue

        if ":" in text and not in_constraints:
            label, _, value = text.partition(":")
            key = _LABELS.get(label.strip().lower())
            if key:
                value = value.strip()
                if key in _NUMERIC_FIELDS:
                    try:
                        num = _clean_number(value)
                    except ValueError:
                        continue
                    ips[key] = int(num) if key == "time_horizon_years" else num
                else:
                    ips[key] = value
                continue

        if in_constraints:
            cleaned = text.lstrip("•-*‣◦").strip()
            if cleaned:
                ips["constraints"].append(cleaned)

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if not any("asset class" in h for h in header_cells):
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 3 or not cells[0]:
                continue
            try:
                ips["allocation_targets"][cells[0]] = {
                    "min": _clean_number(cells[1]),
                    "max": _clean_number(cells[2]),
                }
            except ValueError:
                continue

    if not ips["allocation_targets"]:
        raise IpsDocxParseError(
            "Couldn't find a 'Target Asset Allocation Bands' table with an 'Asset Class' "
            "header -- check the document follows the template's structure."
        )

    return ips
